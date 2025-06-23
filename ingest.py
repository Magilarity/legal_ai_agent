import os
from pathlib import Path

import docx  # python-docx
import fitz  # PyMuPDF
from sqlalchemy import text

from db.schema import ENGINE, Document, Session


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            text_chunks = []
            with fitz.open(path) as pdf:
                for page in pdf:
                    text_chunks.append(page.get_text())
            return "\n".join(text_chunks)
        elif ext in (".docx",):
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            # для інших форматів можна додати логіки, поки – пустий рядок
            return ""
    except Exception as e:
        print(f"[warn] помилка витягу тексту з {path}: {e}")
        return ""


def ingest_all(downloads_folder="downloads"):
    session = Session()
    for tender_dir in Path(downloads_folder).iterdir():
        if not tender_dir.is_dir():
            continue
        tender_id = tender_dir.name
        for file in tender_dir.rglob("*"):
            if not file.is_file():
                continue
            content = extract_text(str(file))
            if content.strip():
                doc = Document(tender=tender_id, title=file.name, content=content)
                session.add(doc)
    session.commit()

    # rebuild FTS index
    with ENGINE.connect() as conn:
        conn.execute(text("INSERT INTO documents_fts(documents_fts) VALUES('rebuild')"))

    print("✅ Ingestion complete.")


if __name__ == "__main__":
    ingest_all()
